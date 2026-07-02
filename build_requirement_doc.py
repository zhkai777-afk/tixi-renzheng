from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "申城体系认证平台需求规格说明书-功能模块版-含列表页.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 116, 139)
BLACK = RGBColor(0, 0, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size=11, color=BLACK, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_para(doc, text="", size=11, color=BLACK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    run = p.add_run(text)
    color = BLUE if level in (1, 2) else DARK_BLUE
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_kv_table(doc, rows, widths=(1800, 7560)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        cells[0].text = ""
        cells[1].text = ""
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=10.5, bold=True, color=DARK_BLUE)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, title in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            run = p.add_run(value)
            set_run_font(run, size=9.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("申城体系认证平台 | 需求规格说明书")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("内部评审稿")
    set_run_font(r, size=9, color=MUTED)
    return doc


def add_cover(doc):
    add_para(doc, "需求规格说明书", size=24, bold=True, after=4)
    add_para(doc, "申城体系认证平台 Demo 功能模块版", size=14, color=MUTED, after=16)
    add_kv_table(doc, [
        ("文档版本", "V0.1 评审稿"),
        ("输出日期", "2026-07-01"),
        ("适用范围", "归档材料提交、初次评定、专家评定、认证决定"),
        ("原型依据", "归档材料提交列表/详情、归档材料退回详情、初次评定列表/详情、专家评定列表/专家评审详情、认证决定列表/详情、业务状态机和业务流程图"),
        ("核心口径", "认证领域统一使用 Q/E/S/F；材料清单统一使用 MSF11-20 审核材料归档/上报清单 48 项。"),
    ])
    add_callout(
        doc,
        "编写原则",
        "本说明书按开发落地视角组织：先说明业务对象、状态流转和数据关系，再按四个业务模块展开功能规格。每个功能模块均覆盖角色、目标、入口、数据、操作、状态、校验、异常、日志、权限和验收标准，减少开发阶段反复追问。"
    )


def add_overview(doc):
    add_heading(doc, "1. 文档说明", 1)
    add_para(doc, "本文档用于将当前申城体系认证 Demo 原型转化为开发可执行的需求规格说明。页面样式、交互细节和模拟数据以现有原型为参考，业务规则以项目需求摘要、MSF11-20 材料清单和业务状态机为准。")
    add_heading(doc, "1.1 本期模块范围", 2)
    add_table(doc, ["模块", "主要角色", "对应页面", "业务目标"], [
        ("归档材料提交", "项目管理人员", "归档材料提交详情、归档材料退回处理详情", "检查系统带入材料、补齐缺失材料，并提交进入初次评定；处理初评或认证决定退回后的补正。"),
        ("初次评定", "初评协调员", "初次评定详情", "接收归档材料，指派专家，汇总专家反馈和问题复核结果，决定通过初评或退回补正。"),
        ("专家评定", "评定专家", "专家评定列表、专家评审详情", "按授权体系审阅材料文件，预览证据，登记问题、补充材料或补充审核要求，并提交个人评定结论。"),
        ("认证决定", "认证决定协调员", "二次评定/认证决定详情", "在初评通过后开展认证决定复核，进行专家覆盖校验、问题清理和体系级通过/退回分流。"),
    ], [1600, 1500, 2600, 3660])

    add_heading(doc, "1.2 统一术语", 2)
    add_table(doc, ["术语", "说明"], [
        ("Q", "质量管理体系。"),
        ("E", "环境管理体系。"),
        ("S", "职业健康安全。"),
        ("F", "食品安全体系。"),
        ("MSF11-20", "审核材料归档/上报清单，当前版本共 48 项材料，是材料归档、评定和问题登记的权威清单。"),
        ("材料清单项", "MSF11-20 中的一项材料，是完整性校验和评定问题登记的最小业务对象。"),
        ("具体文件", "材料清单项下的附件、在线表单或资料库关联文件，仅作为证据展示，不单独承载评定结论。"),
        ("有效问题", "尚未删除、尚未关闭、会影响通过/退回判断的问题记录。"),
    ], [1900, 7460])

    add_heading(doc, "2. 总体业务规则", 1)
    rules = [
        "模拟项目认证领域统一为 Q/E/S/F，不再使用 IT 作为默认体系。",
        "材料清单统一使用 MSF11-20 的 48 项材料，不得在开发实现中写死 44 项、29 项等不一致数量。",
        "评定对象是材料清单项，不是具体文件；具体文件仅作为材料项下的证据。",
        "初次评定、专家评定和认证决定登记问题时，问题必须关联材料清单项，并可关联涉及体系、问题类型、问题描述、处理要求和提交人。",
        "材料明细表固定为“具体文件名称 / 适用体系 / 操作”三列；文件来源、上传时间、类型等放在文件名下方辅助信息中。",
        "主操作按钮禁用时必须展示明确原因，例如缺失材料、专家覆盖不足、存在未关闭问题或当前状态不允许。",
        "初评退回后，项目管理人员补充材料再次提交进入初次评定；认证决定退回后，再次提交直接回到认证决定，不再经过初评。",
        "认证决定通过后，任务进入已完结，相关页面只读，保留预览和历史记录查看能力。",
    ]
    for item in rules:
        add_bullet(doc, item)

    add_heading(doc, "2.1 状态流转", 2)
    add_table(doc, ["状态", "进入条件", "可执行动作", "下一状态"], [
        ("草稿", "系统自动生成归档材料任务，材料未齐全。", "查看材料、补充材料、资料库选择、删除补充文件。", "待提交。"),
        ("待提交", "必交材料在项目涉及体系下均已齐全。", "提交初次评定。", "初次评定。"),
        ("初次评定", "项目管理人员提交归档材料后进入。", "初评协调员指派专家、登记/复核问题、通过或退回。", "二次评定/认证决定，或初次评定退回。"),
        ("初次评定退回", "初评协调员确认退回。", "项目管理人员查看原因并补充材料。", "再次提交后回到初次评定。"),
        ("二次评定/认证决定", "初评通过后进入。", "认证决定协调员指派专家、处理问题、体系级通过或退回。", "已完结，或二次评定退回。"),
        ("二次评定退回", "认证决定退回一个或多个体系。", "项目管理人员按退回原因补充材料。", "再次提交后直接回到二次评定/认证决定。"),
        ("已完结", "认证决定全部通过。", "查看详情、预览文件、查看历史。", "终态，不再流转。"),
    ], [1400, 2600, 3000, 2360])

    add_heading(doc, "2.2 核心数据关系", 2)
    relations = [
        "认证项目 1:N 归档材料任务；本期 Demo 默认一个认证项目生成一个归档材料任务。",
        "归档材料任务 N:M 认证体系；当前任务默认包含 Q/E/S/F 四个体系。",
        "归档材料任务 1:N 材料清单项；清单项来自 MSF11-20，按通用材料、一阶段材料、二阶段及其他审核材料分组。",
        "材料清单项 1:N 具体文件；一个文件可覆盖一个或多个体系。",
        "材料清单项 1:N 评定问题；问题记录可涉及一个或多个体系。",
        "评定专家 N:M 体系/专业代码；专家材料范围和问题登记范围必须受授权限制。",
        "认证决定支持体系级分流；同一任务中部分体系可通过，部分体系可退回。",
    ]
    for item in relations:
        add_bullet(doc, item)


def add_description_template(doc):
    add_heading(doc, "3. 功能描述统一模板", 1)
    add_para(doc, "后续每个功能模块均按以下维度描述。开发拆解任务、接口设计、测试用例和验收时应保持同一口径。")
    add_table(doc, ["维度", "需要说明的内容"], [
        ("功能名称", "使用业务语言命名，例如“提交初次评定”“登记材料问题”“按体系通过认证决定”。"),
        ("使用角色", "明确角色、权限范围和授权体系限制。"),
        ("业务目标", "说明用户进入功能后要完成的判断或动作。"),
        ("入口与前置条件", "入口页面、按钮位置、当前任务状态、必需数据。"),
        ("页面/弹窗要素", "展示区域、字段、按钮、筛选项、表格列和空状态。"),
        ("数据来源", "来自项目、材料清单、文件库、专家库、问题记录、流程状态或操作日志。"),
        ("操作流程", "用户操作步骤、系统响应和页面刷新逻辑。"),
        ("按钮条件", "启用、禁用、只读、隐藏和禁用原因。"),
        ("校验规则", "必填、格式、范围、重复、体系覆盖、状态冲突等。"),
        ("状态流转", "操作成功后任务状态、问题状态、专家状态或体系状态如何变化。"),
        ("接口与保存", "建议接口、入参、出参、落库字段和幂等要求。"),
        ("日志与通知", "操作人、操作时间、原因、对象、前后状态，以及是否触发通知。"),
        ("异常场景", "加载失败、权限不足、并发变更、文件失败、状态已变化等。"),
        ("验收标准", "面向测试的可验证结果。"),
    ], [1900, 7460])


def module_section(doc, num, title, subtitle, meta_rows, common_rules, functions):
    doc.add_page_break()
    add_heading(doc, f"{num}. {title}", 1)
    add_para(doc, subtitle)
    add_heading(doc, f"{num}.1 模块定位", 2)
    add_kv_table(doc, meta_rows)
    add_heading(doc, f"{num}.2 模块通用规则", 2)
    for item in common_rules:
        add_bullet(doc, item)
    for idx, fn in enumerate(functions, start=3):
        add_heading(doc, f"{num}.{idx} {fn['name']}", 2)
        add_kv_table(doc, [
            ("使用角色", fn["role"]),
            ("业务目标", fn["goal"]),
            ("入口与前置条件", fn["entry"]),
            ("数据来源", fn["data"]),
        ])
        add_heading(doc, "页面/弹窗要素", 3)
        for item in fn["elements"]:
            add_bullet(doc, item)
        add_heading(doc, "操作流程", 3)
        for item in fn["flow"]:
            add_number(doc, item)
        add_heading(doc, "校验、状态与权限", 3)
        for item in fn["rules"]:
            add_bullet(doc, item)
        add_heading(doc, "接口与保存建议", 3)
        for item in fn["api"]:
            add_bullet(doc, item)
        add_heading(doc, "异常与验收标准", 3)
        for item in fn["accept"]:
            add_bullet(doc, item)


def add_modules(doc):
    module_section(
        doc,
        4,
        "归档材料提交（项目管理人员）",
        "该模块由项目管理人员使用，核心目标是确认系统带入的前序材料是否满足 MSF11-20 归档要求，补齐缺失材料，并按当前流程提交到初次评定或认证决定。",
        [
            ("对应原型", "归档材料提交-详情-10体系重构版.html；归档材料退回-详情-简化布局-v2-顶部统计卡片版.html"),
            ("使用角色", "项目管理人员"),
            ("关键目标", "补齐材料、判断是否可提交、处理退回原因、重新提交。"),
            ("核心页面结构", "顶部任务信息 + 左侧材料导航 + 中间材料台账/明细 + 右侧完整性控制台。"),
            ("主要状态", "草稿、待提交、初次评定退回、二次评定退回。"),
        ],
        [
            "左侧材料导航必须展示 MSF11-20 48 项材料，按通用材料、一阶段材料、二阶段及其他审核材料分组。",
            "第一项为“全部材料跨体系视图”，与普通材料项使用一致样式。",
            "快速筛选建议为“全部 / 缺失 / 已带入”；筛选结果同时影响左侧材料树和中间展示区域。",
            "材料完整性按项目涉及体系分别计算，所有必交材料齐全后主按钮才启用。",
            "退回场景必须展示退回来源、退回原因、涉及体系、涉及材料和补正要求。",
        ],
        [
            {
                "name": "归档材料提交列表页",
                "role": "项目管理人员、归档材料管理人员。",
                "goal": "按任务状态集中查看项目归档材料任务，快速识别待提交、退回补正和已完成任务，并进入详情处理。",
                "entry": "从菜单“材料归档管理”进入；当前用户仅能查看其权限范围内的归档任务。",
                "data": "归档材料任务、项目客户、认证领域、认证类型、归属机构、归属部门、计划审核时间、材料进度、任务状态。",
                "elements": [
                    "页面标题为“材料归档管理”，顶部可保留一段功能描述占位。",
                    "状态页签按业务状态分组：全部、草稿、待提交、已提交初评、已提交认证决定、初次评定退回、认证决定退回、已完结；页签可显示待处理数量角标。",
                    "查询区包含任务号、客户名称、状态，支持“查询”“重置”“展开/收起高级条件”。",
                    "类型筛选包含全部类型、初次认证、监督审核、再认证等认证类型标签。",
                    "列表列建议为任务号、客户名称、认证领域、认证类型、归属机构、归属部门、计划审核时间、材料进度、状态、操作。",
                    "材料进度以进度条和数量展示，例如 34/40；进度异常或草稿状态可用风险色提示。",
                    "操作列根据状态展示“提交”“详情”等入口；不可操作时只保留详情。",
                    "底部提供每页条数、页码、上一页/下一页和跳页能力。",
                ],
                "flow": [
                    "用户进入列表页后，系统默认查询全部权限范围内任务，并按最近更新时间或计划审核时间倒序展示。",
                    "用户点击状态页签，系统按状态快速过滤列表并刷新数量。",
                    "用户输入任务号、客户名称或状态后点击查询，系统返回符合条件的任务。",
                    "用户点击任务号或“详情”进入归档材料提交详情页。",
                    "对待提交任务，用户可从列表操作列直接点击“提交”，系统仍需执行完整性校验和确认弹窗。",
                ],
                "rules": [
                    "页签状态与归档任务状态机一致，不得出现无法流转的中间状态。",
                    "列表页“提交”仅在任务达到待提交条件且当前用户有权限时显示或启用。",
                    "材料进度必须来自完整性校验结果，不允许前端按文件数量自行推断。",
                    "认证领域统一使用 Q/E/S/F 徽标展示，不使用 IT。",
                    "分页、查询和页签切换需要保留当前筛选上下文，返回列表时应尽量恢复上次查询条件。",
                ],
                "api": [
                    "GET /archive-tasks：分页查询归档材料任务列表。",
                    "查询参数建议包含 taskNo、customerName、status、certificationType、organizationId、departmentId、pageNo、pageSize。",
                    "返回字段建议包含 taskId、taskNo、customerName、systems、certificationTypes、organizationName、departmentName、plannedAuditDate、materialDoneCount、materialTotalCount、status、availableActions。",
                    "POST /archive-tasks/{taskId}/submit：支持从列表触发提交，但提交前后端必须复用详情页校验逻辑。",
                ],
                "accept": [
                    "状态页签切换后列表状态与页签一致，数量角标准确。",
                    "待提交任务显示“提交 | 详情”，其他状态按权限显示详情或补正入口。",
                    "材料进度与详情页右侧完整性统计一致。",
                    "查询、重置、展开高级条件、分页跳转均可正常使用。",
                ],
            },
            {
                "name": "材料清单查看与筛选",
                "role": "项目管理人员。",
                "goal": "快速定位缺失、已带入或某体系相关材料，判断当前任务是否具备提交条件。",
                "entry": "进入归档材料提交或退回处理详情页；任务未完结且当前角色具备项目材料管理权限。",
                "data": "项目基本信息、项目涉及体系、MSF11-20 材料清单、材料文件、完整性统计、退回问题。",
                "elements": [
                    "顶部展示任务号、客户、项目体系 Q/E/S/F、当前状态和主操作按钮。",
                    "左侧提供认证领域多选、认证类型筛选、关键字搜索、快速筛选和分组材料树。",
                    "材料项第一行显示材料名称，第二行显示体系徽标，底部显示分类/必交/文件数或缺失状态。",
                    "中间区域支持材料台账视图和具体文件明细视图。",
                    "右侧控制台展示总体材料项数、已补齐、缺失项、各体系完整度和阻断原因。",
                ],
                "flow": [
                    "用户进入页面后，系统加载项目、体系范围和 MSF11-20 清单。",
                    "系统聚合前序环节文件和人工补充文件，计算每项材料覆盖体系。",
                    "用户通过体系、状态或关键词筛选材料。",
                    "点击材料项后，中间区域切换到该材料的具体文件列表或缺失空状态。",
                    "筛选变化后，右侧完整性统计同步刷新。",
                ],
                "rules": [
                    "筛选不应删除材料清单，只改变当前展示范围。",
                    "无文件材料必须显示缺失/未上传，不得从列表中移除。",
                    "已完结状态下材料清单只读，保留预览能力。",
                    "认证决定退回补正时，只允许编辑需补正材料或仍可按权限补充其他材料，具体范围由业务配置控制。",
                ],
                "api": [
                    "GET /archive-tasks/{taskId}：获取任务基础信息和当前状态。",
                    "GET /archive-tasks/{taskId}/materials：获取材料清单、文件、完整性和退回标记。",
                    "GET /archive-tasks/{taskId}/completion：获取各体系完整性和阻断原因。",
                ],
                "accept": [
                    "默认展示 48 项材料，且分组数量与 MSF11-20 一致。",
                    "切换体系筛选后，体系徽标和统计同步变化。",
                    "点击“仅看缺失”时，左侧只显示缺失材料，中间区域显示缺失说明和补充入口。",
                    "页面不得出现 IT 作为默认体系。",
                ],
            },
            {
                "name": "本地上传与资料库选择",
                "role": "项目管理人员。",
                "goal": "对缺失或需补正材料进行文件补充，并维护文件覆盖体系。",
                "entry": "材料项详情、材料台账行或缺失空状态中的“本地上传”“资料库选择”按钮。",
                "data": "当前材料项、适用体系、缺失体系、资料库文件、上传文件元数据。",
                "elements": [
                    "上传面板默认选中当前材料尚未覆盖的体系。",
                    "资料库弹窗支持分类、关键词、体系、文件类型和环节筛选。",
                    "文件行展示具体文件名称、所属材料、来源、上传时间、适用体系和预览/删除操作。",
                    "删除补充文件必须使用自定义确认弹窗，并说明删除后将重新计算完整性。",
                ],
                "flow": [
                    "用户选择当前材料项并点击补充入口。",
                    "系统打开上传面板或资料库弹窗，自动带出材料名称和建议体系。",
                    "用户选择文件和覆盖体系后确认。",
                    "系统保存文件关联关系，并重新计算材料覆盖和任务完整性。",
                    "页面刷新材料项状态、文件数、右侧完整性和主按钮状态。",
                ],
                "rules": [
                    "上传或资料库选择时，至少选择一个覆盖体系。",
                    "覆盖体系必须属于当前材料适用体系和项目涉及体系交集。",
                    "同一材料下重复文件应提示用户确认或按业务规则拦截。",
                    "删除文件后如导致必交材料缺失，主按钮必须立即禁用并展示原因。",
                ],
                "api": [
                    "POST /archive-tasks/{taskId}/materials/{materialId}/files：上传或关联资料库文件。",
                    "DELETE /archive-tasks/{taskId}/materials/{materialId}/files/{fileId}：删除人工补充关联。",
                    "GET /library/files：查询资料库候选文件。",
                    "保存字段包括 materialId、fileId、fileName、sourceType、coveredSystems、uploadedBy、uploadedAt。",
                ],
                "accept": [
                    "上传成功后材料项文件数增加，缺失体系减少。",
                    "资料库选择多个文件后可一次性关联到当前材料。",
                    "删除文件后完整性重新计算，且操作日志可追溯。",
                    "所有反馈均使用页面内 toast 或 modal，不使用浏览器默认 alert/confirm。",
                ],
            },
            {
                "name": "提交初次评定/重新提交",
                "role": "项目管理人员。",
                "goal": "在材料完整性满足要求后，将任务提交到下一评定节点。",
                "entry": "顶部或右侧控制台主按钮；草稿/待提交/退回补正状态下可见。",
                "data": "任务状态、完整性校验结果、退回来源、补正记录、当前用户权限。",
                "elements": [
                    "主按钮文案在首次提交时为“提交初次评定”。",
                    "初评退回后再次提交仍进入初次评定，可沿用“重新提交初次评定”或业务确认后的文案。",
                    "认证决定退回后再次提交应进入认证决定，按钮文案需体现“重新提交认证决定”。",
                    "提交确认弹窗展示即将进入的下一节点、材料完整性和未解决风险提示。",
                ],
                "flow": [
                    "用户点击提交按钮。",
                    "系统执行完整性校验和状态校验。",
                    "校验通过后打开确认弹窗。",
                    "用户确认提交后，系统变更任务状态并写入流程日志。",
                    "页面显示提交成功，并跳转或提示进入对应评定节点。",
                ],
                "rules": [
                    "存在必交材料缺失时按钮禁用，必须展示缺失体系和缺失项数。",
                    "任务已进入评定中或已完结时，不允许重复提交。",
                    "认证决定退回后的重新提交不得回到初评。",
                    "提交接口需要防重复点击和幂等处理。",
                ],
                "api": [
                    "POST /archive-tasks/{taskId}/submit：提交或重新提交任务。",
                    "请求需包含 currentStatus、submitTarget、补正说明可选字段。",
                    "响应返回新状态、下一处理角色、流程实例节点和提示文案。",
                ],
                "accept": [
                    "材料缺失时提交按钮不可点击，且原因明确。",
                    "材料齐全时提交按钮启用，确认后状态正确流转。",
                    "初评退回和认证决定退回的重新提交路径符合状态机。",
                    "提交成功后产生流程日志和通知待办。",
                ],
            },
        ],
    )

    module_section(
        doc,
        5,
        "初次评定（初评协调员）",
        "该模块由初评协调员使用，核心目标是接收归档材料、完成专家覆盖、汇总专家意见和整改问题，并决定任务是否通过初评或退回项目管理人员补正。",
        [
            ("对应原型", "初次评定-详情-复核 2.html"),
            ("使用角色", "初评协调员/初评管理员"),
            ("关键目标", "指派专家、复核问题、判断通过/退回条件。"),
            ("核心页面结构", "顶部任务信息 + 左侧评审材料清单 + 中间问题/文件工作区 + 右侧初评控制台。"),
            ("主要状态", "初次评定、初次评定退回、二次评定/认证决定。"),
        ],
        [
            "初评控制台必须展示专家覆盖、专家提交进度、有效问题、待复核问题和当前可执行动作。",
            "问题登记和复核均落在材料清单项级别，不对具体文件逐行给出评定状态。",
            "专家覆盖不足时可以保存专家指派进度，但不得通过初评。",
            "存在有效未关闭问题时不得通过初评，可执行退回任务。",
        ],
        [
            {
                "name": "初次评定列表页",
                "role": "初评协调员/初评管理员。",
                "goal": "集中接收归档材料初评任务，识别待接收、待指派专家、专家评审中、待初评确认、初评退回和初评通过任务，并进入详情处理。",
                "entry": "从菜单“初次评定管理”进入；仅展示流转到初评节点或与当前协调员相关的任务。",
                "data": "初评任务、归档任务、客户、认证领域、认证类型、归属机构、归属部门、任务提交时间、初评状态、专家指派状态。",
                "elements": [
                    "页面标题为“初次评定管理”。",
                    "状态页签包含全部、待接收、待指派专家、专家评审中、待初评确认、初评退回、初评通过；待接收页签可显示数量角标。",
                    "查询区包含任务号、客户名称、状态，支持查询、重置和展开高级条件。",
                    "认证类型标签包含全部类型、初次认证、监督审核、再认证。",
                    "列表列建议为任务号、客户名称、认证领域、认证类型、归属机构、归属部门、任务提交时间、状态、操作。",
                    "操作列在待接收状态显示“接收 | 详情”，其他状态默认显示“详情”，必要时可显示“处理”。",
                    "状态标签需覆盖待接收、待指派专家、专家评审中、初评待确认、初评退回、初评通过。",
                    "底部提供分页、每页条数和跳页。",
                ],
                "flow": [
                    "系统默认展示当前协调员可处理的全部初评任务。",
                    "协调员点击待接收页签，快速查看新提交任务。",
                    "点击“接收”后系统将任务领取到当前协调员名下，并进入或刷新详情页。",
                    "点击“详情”进入初次评定详情页，进行专家指派、问题复核和通过/退回。",
                    "任务在详情页状态变化后，返回列表时应刷新当前行状态。",
                ],
                "rules": [
                    "未流转到初评节点的归档任务不得出现在初评列表。",
                    "待接收任务可由具备权限的协调员接收；已接收任务应显示处理人或按权限限制操作。",
                    "初评通过任务默认只读，保留详情查看。",
                    "初评退回任务在初评列表可查看，但补正操作由项目管理人员在归档材料模块处理。",
                ],
                "api": [
                    "GET /first-review/tasks：分页查询初评任务列表。",
                    "查询参数建议包含 taskNo、customerName、status、certificationType、submittedAtRange、pageNo、pageSize。",
                    "POST /first-review/tasks/{taskId}/receive：接收初评任务。",
                    "返回字段建议包含 taskId、taskNo、customerName、systems、certificationTypes、organizationName、departmentName、submittedAt、status、assigneeName、availableActions。",
                ],
                "accept": [
                    "待接收任务可在列表直接接收，接收后状态变化或进入详情。",
                    "不同页签的状态过滤准确，任务提交时间展示到秒。",
                    "初评退回和初评通过状态能在列表中被正确识别。",
                    "列表任务进入详情后，详情页角色和任务状态与列表一致。",
                ],
            },
            {
                "name": "专家指派与覆盖校验",
                "role": "初评协调员。",
                "goal": "为任务分配覆盖 Q/E/S/F 体系的专家，并校验授权体系覆盖是否满足评定要求。",
                "entry": "初评详情页右侧“专家覆盖”卡片中的“指派/调整”按钮。",
                "data": "专家库、专家授权体系、专业代码、当前指派记录、项目体系。",
                "elements": [
                    "专家弹窗提供领域、姓名、机构等筛选。",
                    "候选专家列表展示姓名、所属机构、授权体系、专业代码、当前负责项目数和预计完成时间。",
                    "覆盖预览实时展示已覆盖体系和缺失体系。",
                    "已指派专家卡片展示负责体系、提交状态、问题数和调整入口。",
                ],
                "flow": [
                    "协调员打开专家指派弹窗。",
                    "系统加载候选专家并排除已指派专家。",
                    "协调员选择专家及其负责体系。",
                    "系统实时计算覆盖预览。",
                    "确认后保存指派记录并通知专家。",
                ],
                "rules": [
                    "专家负责体系必须在其授权范围内。",
                    "同一体系可由多个专家覆盖，但通过初评前所有项目体系必须至少被覆盖。",
                    "替换专家时必须填写原因，并保留原专家相关问题和处理历史。",
                    "专家已提交后调整专家需记录日志，并按业务规则决定是否重新评审。",
                ],
                "api": [
                    "GET /experts：查询专家候选列表。",
                    "POST /first-review/{taskId}/experts：保存专家指派。",
                    "POST /first-review/{taskId}/experts/{assignmentId}/replace：替换专家并保存原因。",
                ],
                "accept": [
                    "未覆盖全部体系时，控制台显示缺失体系。",
                    "指派成功后专家卡片出现在右侧控制台。",
                    "替换专家后保留调整日志。",
                    "专家覆盖不足时“通过初评”不可用。",
                ],
            },
            {
                "name": "问题登记、复核与历史",
                "role": "初评协调员。",
                "goal": "登记协调员发现的问题，复核专家或项目管理人员处理结果，形成通过/退回判断依据。",
                "entry": "中间问题清单、材料详情问题区或右侧控制台问题统计入口。",
                "data": "材料清单项、问题记录、整改说明、整改附件、操作日志、专家反馈。",
                "elements": [
                    "问题清单展示材料、涉及体系、问题类型、问题描述、来源、复核状态和操作。",
                    "问题类型至少包括材料问题、补充材料、补充审核要求。",
                    "问题登记弹窗包含材料项、问题类型、涉及体系、问题描述、处理要求和附件。",
                    "历史弹窗展示问题创建、修改、通过、退回、撤回和删除记录。",
                ],
                "flow": [
                    "协调员在全部材料或单项材料范围内打开问题登记弹窗。",
                    "选择材料项和问题类型，系统按材料适用体系生成可选体系。",
                    "填写问题描述和处理要求后保存。",
                    "项目管理人员补正后，协调员执行通过、退回或撤回操作。",
                    "系统更新问题状态，并刷新控制台判断条件。",
                ],
                "rules": [
                    "问题描述必填，涉及体系至少选择一个。",
                    "问题不能关联到具体文件作为评定结论对象。",
                    "删除问题、退回复核、撤回通过必须填写原因。",
                    "已关闭问题不再阻断通过，但仍需保留历史。",
                ],
                "api": [
                    "POST /review-issues：新增问题。",
                    "PUT /review-issues/{issueId}：编辑问题。",
                    "POST /review-issues/{issueId}/pass：复核通过。",
                    "POST /review-issues/{issueId}/return：退回复核。",
                    "DELETE /review-issues/{issueId}：逻辑删除问题并保存原因。",
                ],
                "accept": [
                    "问题保存后在对应材料项和全部问题清单中同步出现。",
                    "问题状态变化后控制台有效问题数实时更新。",
                    "问题历史可按单条问题查看。",
                    "无有效问题且专家覆盖/提交满足条件时，才允许通过初评。",
                ],
            },
            {
                "name": "通过初评/退回任务",
                "role": "初评协调员。",
                "goal": "基于专家覆盖、专家评审完成情况和问题关闭情况，给出初评结论。",
                "entry": "顶部操作区或右侧“整改复核判断”卡片。",
                "data": "专家覆盖、专家提交状态、有效问题、材料文件、任务状态。",
                "elements": [
                    "通过初评按钮显示在满足条件时启用。",
                    "退回任务按钮在存在有效问题或需要项目管理人员补正时启用。",
                    "确认弹窗展示通过或退回后的下一节点、阻断项和处理说明。",
                    "退回时必须填写退回原因，可关联具体问题和材料项。",
                ],
                "flow": [
                    "协调员查看右侧控制台判断结果。",
                    "满足通过条件时点击通过初评并确认。",
                    "存在阻断问题时点击退回任务并填写退回原因。",
                    "系统变更任务状态，生成流程日志和待办通知。",
                ],
                "rules": [
                    "通过初评条件：专家覆盖完整、必要专家已提交、无有效未关闭问题、任务仍处于初评中。",
                    "退回条件：存在有效问题、补正要求或初评协调员确认材料不满足要求。",
                    "任务状态变化后，初评页面进入只读或结果态。",
                ],
                "api": [
                    "POST /first-review/{taskId}/pass：通过初评。",
                    "POST /first-review/{taskId}/return：退回归档材料补正。",
                    "接口需校验当前状态和版本号，避免并发重复决策。",
                ],
                "accept": [
                    "通过后任务进入认证决定节点。",
                    "退回后任务进入初次评定退回，项目管理人员可查看退回原因。",
                    "不满足条件时按钮禁用并展示明确阻断原因。",
                    "所有决策均写入流程日志。",
                ],
            },
        ],
    )

    module_section(
        doc,
        6,
        "专家评定（专家）",
        "该模块由评定专家使用，核心目标是按授权体系审阅材料和文件，登记材料项级别的问题、补充材料或补充审核要求，并提交个人评定结果给协调员。",
        [
            ("对应原型", "专家评审详情-预览登记优化 2.html；专家评定管理列表页截图"),
            ("使用角色", "评定专家"),
            ("关键目标", "按授权体系审阅、预览文件、登记问题、提交评定。"),
            ("核心页面结构", "顶部任务信息 + 左侧授权材料清单 + 中间问题与材料明细 + 右侧专家任务状态。"),
            ("主要状态", "待评审、评审中、已提交、被退回复核/需补充说明。"),
        ],
        [
            "专家只展示包含其授权体系的材料；完全不包含授权体系的材料不进入专家材料树。",
            "专家可登记的问题必须关联材料清单项，并限制在其授权体系范围内。",
            "专家可预览文件并在预览窗口直接发起问题登记，但问题仍归属材料项。",
            "专家提交后，除非协调员退回或允许补充，否则专家侧进入只读。",
        ],
        [
            {
                "name": "专家评定列表页",
                "role": "评定专家。",
                "goal": "让专家集中查看分配给自己的评定任务，识别待评定、进行中、已登记问题、待复核、已退回和已完成任务，并进入详情开展材料审阅。",
                "entry": "专家从菜单“专家评定管理”或个人待办进入；仅展示当前专家被指派的任务。",
                "data": "专家指派记录、任务基础信息、客户、认证领域、认证类型、归属机构、归属部门、任务截止日期、专家评定状态、问题统计。",
                "elements": [
                    "页面标题可归入“专家评定管理”，如沿用初评管理框架，需在标题或菜单上明确当前为专家评定。",
                    "状态页签包含全部、待评定、已退回、进行中、已完成；根据业务需要可区分待复核。",
                    "查询区包含任务号、客户名称、状态，支持查询、重置和展开高级条件。",
                    "认证类型标签包含全部类型、初次认证、监督审核、再认证。",
                    "列表列建议为任务号、客户名称、认证领域、认证类型、归属机构、归属部门、任务截止日期、状态、操作。",
                    "状态标签建议覆盖待评审/待评定、评审中/评定中、无有效问题、已登记问题、待复核、已退回、已完成。",
                    "操作列默认显示“详情”；待评定或进行中任务点击详情进入专家评审详情页。",
                    "底部提供分页、每页条数和跳页。",
                ],
                "flow": [
                    "专家进入列表页后，系统按当前登录专家过滤任务。",
                    "专家通过页签定位待评定或进行中任务。",
                    "专家可按任务号、客户名称、状态查询。",
                    "点击详情进入专家评审详情页，查看授权材料、预览文件并登记问题。",
                    "专家提交评定后，返回列表时任务状态更新为已完成或待协调员确认的结果态。",
                ],
                "rules": [
                    "专家列表不得展示未指派给当前专家的任务。",
                    "专家只能看到其授权体系相关的任务信息和材料范围。",
                    "任务截止日期用于提醒，不直接替代流程状态。",
                    "已完成任务默认只读，若协调员退回复核，则重新进入待评定或已退回状态。",
                ],
                "api": [
                    "GET /expert-review/tasks：分页查询当前专家评定任务。",
                    "查询参数建议包含 taskNo、customerName、status、certificationType、deadlineRange、pageNo、pageSize。",
                    "返回字段建议包含 assignmentId、taskId、taskNo、customerName、systems、certificationTypes、organizationName、departmentName、deadline、status、issueCount、availableActions。",
                ],
                "accept": [
                    "当前专家只能看到自己的评定任务。",
                    "待评定、进行中、已完成等页签过滤准确。",
                    "点击详情后进入专家评审详情页，材料范围受授权体系限制。",
                    "专家提交后列表状态同步更新。",
                ],
            },
            {
                "name": "授权材料范围与筛选",
                "role": "评定专家。",
                "goal": "确保专家只处理被授权的体系材料，降低误评其他体系材料的风险。",
                "entry": "专家从待办进入专家评审详情页。",
                "data": "专家指派记录、授权体系、MSF11-20 材料清单、材料文件、问题记录。",
                "elements": [
                    "左侧显示专家姓名、角色、授权体系和当前可评审材料数。",
                    "体系筛选在单体系专家场景下可固定为授权体系；多体系专家可在授权范围内切换。",
                    "快速筛选包括全部、有问题、无问题、无文件。",
                    "材料树统计按授权后的材料范围计算。",
                ],
                "flow": [
                    "专家进入页面后，系统读取其指派和授权体系。",
                    "系统过滤材料清单，只保留包含授权体系的材料。",
                    "专家通过筛选或搜索定位材料。",
                    "点击材料项后，中间区域展示该材料的问题和文件证据。",
                ],
                "rules": [
                    "专家不得查看与授权体系完全无关的材料。",
                    "跨体系材料可展示全部体系徽标，但非授权体系应只读或弱化。",
                    "无文件材料仍需展示，以便专家登记补充材料要求。",
                ],
                "api": [
                    "GET /expert-review/{assignmentId}：获取专家评审任务。",
                    "GET /expert-review/{assignmentId}/materials：获取授权过滤后的材料和文件。",
                ],
                "accept": [
                    "Q 专家只看到包含 Q 的材料。",
                    "统计数量与过滤后的材料范围一致。",
                    "搜索和快速筛选不会越权展示材料。",
                ],
            },
            {
                "name": "文件预览与问题登记",
                "role": "评定专家。",
                "goal": "在查看具体证据时快速登记材料项级别问题。",
                "entry": "材料明细表“预览/查看表单”按钮，或问题面板“登记问题”按钮。",
                "data": "文件信息、材料项、已有问题、在线表单内容、附件预览信息。",
                "elements": [
                    "预览弹窗左侧提供当前评审范围内的材料文件目录。",
                    "预览区域展示附件或在线表单模拟内容。",
                    "预览工具区可切换上一份/下一份文件，并直接登记问题。",
                    "问题登记弹窗包含问题类型、材料项、涉及体系、问题描述和处理要求。",
                ],
                "flow": [
                    "专家点击文件预览。",
                    "系统打开预览弹窗并构建当前授权范围内的文件目录。",
                    "专家在预览中发现问题后点击登记问题。",
                    "系统自动带入当前文件所属材料项，但保存时只关联材料项。",
                    "保存后问题出现在材料问题清单和右侧统计中。",
                ],
                "rules": [
                    "问题涉及体系只能从专家授权体系与材料适用体系交集中选择。",
                    "问题描述必填，问题类型必选。",
                    "从预览发起的问题不得把具体文件作为评定结论对象。",
                    "专家已提交后不得新增或删除问题，除非任务被退回给专家。",
                ],
                "api": [
                    "GET /files/{fileId}/preview：获取预览信息或预览地址。",
                    "POST /expert-review/{assignmentId}/issues：专家登记问题。",
                    "DELETE /expert-review/{assignmentId}/issues/{issueId}：删除专家本人未提交前的问题。",
                ],
                "accept": [
                    "从文件预览登记问题后，问题归属该文件所属材料项。",
                    "专家无法选择非授权体系。",
                    "问题保存后材料树显示问题数。",
                    "预览切换不会丢失未保存的问题草稿，或需在切换前明确提示。",
                ],
            },
            {
                "name": "提交专家评定",
                "role": "评定专家。",
                "goal": "完成个人授权范围内材料审阅后提交个人评定结果，供协调员汇总判断。",
                "entry": "专家评审详情页顶部或右侧任务状态卡片中的“提交评定”按钮。",
                "data": "专家授权材料范围、已查看材料、问题记录、专家评定状态。",
                "elements": [
                    "右侧任务状态展示授权材料数、已查看材料数、问题数和提交状态。",
                    "提交前确认弹窗展示当前登记问题和未查看材料提示。",
                    "提交成功后页面进入只读，保留预览和历史查看。",
                ],
                "flow": [
                    "专家完成材料审阅。",
                    "点击提交评定。",
                    "系统执行提交前校验，并展示确认信息。",
                    "专家确认后，系统更新专家评定状态为已提交。",
                    "协调员侧专家进度和问题清单同步更新。",
                ],
                "rules": [
                    "是否要求全部授权材料已查看后才能提交，由业务配置决定；若不强制，应给出未查看提示。",
                    "提交后专家不能编辑问题，除非协调员退回评审。",
                    "无问题提交也是有效结论，需记录为该专家已提交。",
                ],
                "api": [
                    "POST /expert-review/{assignmentId}/submit：提交专家评定。",
                    "请求需携带当前问题版本号，防止并发覆盖。",
                ],
                "accept": [
                    "提交成功后专家状态变为已提交。",
                    "协调员侧进度实时体现专家已提交。",
                    "专家无问题提交时，系统记录“无问题”结论。",
                    "已提交后编辑入口禁用并显示原因。",
                ],
            },
        ],
    )

    module_section(
        doc,
        7,
        "认证决定（认证决定协调员）",
        "该模块对应当前 Demo 中的二次评定/复评页面。面向业务命名建议使用“认证决定”，由认证决定协调员在初评通过后进行最终复核、专家覆盖、问题清理和体系级通过/退回分流。",
        [
            ("对应原型", "二次评定-详情-体系分流按体系卡片版.html"),
            ("使用角色", "认证决定协调员"),
            ("关键目标", "完成认证决定复核、体系级分流、最终放行或退回补正。"),
            ("核心页面结构", "顶部任务决策条 + 左侧归档材料清单 + 中间问题/材料工作区 + 右侧认证决定控制台。"),
            ("主要状态", "二次评定/认证决定、二次评定退回、已完结。"),
        ],
        [
            "认证决定阶段重点展示专家覆盖、体系级问题统计、可通过体系、需退回体系和问题清理情况。",
            "支持全任务通过、全任务退回，以及按体系部分通过/部分退回。",
            "退回后项目管理人员补正并重新提交时，应直接回到认证决定，不再经过初评。",
            "认证决定通过后任务终态只读，保留材料、问题、专家和流程历史。",
        ],
        [
            {
                "name": "认证决定列表页",
                "role": "认证决定协调员。",
                "goal": "集中管理进入认证决定阶段的任务，识别待评定、进行中、待复核、已退回和已通过任务，并进入详情完成最终决定。",
                "entry": "从菜单“认证决定管理”进入；如系统历史菜单为“二次评定管理”，页面文案应逐步统一为认证决定。",
                "data": "认证决定任务、客户、认证领域、认证类型、归属机构、归属部门、任务提交时间、认证决定状态、体系分流结果。",
                "elements": [
                    "页面标题建议为“认证决定管理”；若保留二次评定标题，需要在业务文案中明确其含义为认证决定。",
                    "状态页签包含全部、待评定、已退回、进行中、已通过；根据流程可扩展待复核。",
                    "查询区包含任务号、客户名称、状态，支持查询、重置和展开高级条件。",
                    "认证类型标签包含全部类型、初次认证、监督审核、再认证。",
                    "列表列建议为任务号、客户名称、认证领域、认证类型、归属机构、归属部门、任务提交时间、状态、操作。",
                    "状态标签建议覆盖待评定、进行中、已通过、已退回、待复核。",
                    "操作列默认显示“详情”，进入认证决定详情页处理专家覆盖、问题清理和体系分流。",
                    "底部提供分页、每页条数和跳页。",
                ],
                "flow": [
                    "认证决定协调员进入列表页后，系统展示已初评通过并流转到认证决定阶段的任务。",
                    "协调员通过页签快速查看待处理或已退回任务。",
                    "点击详情进入认证决定详情页。",
                    "在详情页完成通过、退回或体系级分流后，返回列表刷新状态。",
                    "已通过任务进入终态，只允许查看详情和历史。",
                ],
                "rules": [
                    "未初评通过的任务不得进入认证决定列表。",
                    "认证决定退回任务仍出现在列表中，但项目管理人员补正入口在归档材料模块。",
                    "部分体系退回时，列表状态可显示为已退回或待复核，详情页必须展示具体体系分流结果。",
                    "已通过任务对应归档任务状态为已完结。",
                ],
                "api": [
                    "GET /decision-review/tasks：分页查询认证决定任务列表。",
                    "查询参数建议包含 taskNo、customerName、status、certificationType、submittedAtRange、pageNo、pageSize。",
                    "返回字段建议包含 taskId、taskNo、customerName、systems、certificationTypes、organizationName、departmentName、submittedAt、status、systemDecisionSummary、availableActions。",
                ],
                "accept": [
                    "只有进入认证决定节点的任务出现在列表。",
                    "待评定、进行中、已退回、已通过页签过滤准确。",
                    "从列表进入详情后可看到体系级分流和问题清理情况。",
                    "认证决定通过后列表状态为已通过，任务详情只读。",
                ],
            },
            {
                "name": "认证决定专家覆盖",
                "role": "认证决定协调员。",
                "goal": "确认认证决定阶段的专家或复核人员覆盖项目体系和专业代码要求。",
                "entry": "认证决定详情页右侧“专家覆盖”卡片。",
                "data": "专家库、授权体系、专业代码、认证决定阶段指派记录、项目体系。",
                "elements": [
                    "专家弹窗展示领域、姓名、机构、专业代码和分页。",
                    "覆盖校验面板按 Q/E/S/F 展示是否覆盖及缺失代码。",
                    "专家卡片展示负责体系、专业代码、机构和操作入口。",
                ],
                "flow": [
                    "协调员打开专家指派或调整弹窗。",
                    "选择专家及其覆盖体系/代码。",
                    "系统实时计算覆盖结果。",
                    "确认指派后保存并刷新右侧控制台。",
                ],
                "rules": [
                    "通过认证决定前，必须满足业务配置的专家覆盖规则。",
                    "专家覆盖可允许阶段性部分覆盖，但控制台必须提示缺失项。",
                    "删除或调整专家需保留操作日志。",
                ],
                "api": [
                    "GET /decision-review/{taskId}/experts：获取认证决定专家指派。",
                    "POST /decision-review/{taskId}/experts：保存指派。",
                    "DELETE /decision-review/{taskId}/experts/{assignmentId}：删除指派并记录原因。",
                ],
                "accept": [
                    "专家覆盖不足时认证决定通过按钮禁用。",
                    "选择专家专业代码后覆盖预览实时变化。",
                    "专家调整日志在右侧控制台可查看。",
                ],
            },
            {
                "name": "问题清理与体系分流",
                "role": "认证决定协调员。",
                "goal": "按体系识别有效问题，判断哪些体系可通过、哪些体系需退回。",
                "entry": "中间问题清单、体系分流卡片和右侧“体系分流”控制台。",
                "data": "初评问题、专家问题、认证决定问题、问题状态、材料项、涉及体系。",
                "elements": [
                    "问题清单支持按体系、问题类型、状态筛选。",
                    "体系分流卡片按 Q/E/S/F 展示问题数、可通过状态和退回原因。",
                    "问题操作包括登记、编辑、删除、通过、退回、撤回和查看历史。",
                    "每个体系的通过/退回判断应给出可解释原因。",
                ],
                "flow": [
                    "协调员查看各体系有效问题和专家覆盖情况。",
                    "对已整改问题执行通过或退回复核。",
                    "必要时新增认证决定阶段问题。",
                    "系统按体系重新计算可通过/需退回分组。",
                    "协调员执行单体系通过或整体决策。",
                ],
                "rules": [
                    "问题涉及多个体系时，会同时影响多个体系的分流判断。",
                    "存在有效未关闭问题的体系不得通过。",
                    "删除、退回、撤回操作必须填写原因并写入历史。",
                    "体系已通过后，除非业务允许撤回，否则该体系进入只读或结果态。",
                ],
                "api": [
                    "GET /decision-review/{taskId}/issues：获取认证决定问题和初评遗留问题。",
                    "POST /decision-review/{taskId}/issues：新增认证决定问题。",
                    "POST /decision-review/{taskId}/issues/{issueId}/pass：问题通过。",
                    "POST /decision-review/{taskId}/issues/{issueId}/return：问题退回。",
                    "GET /decision-review/{taskId}/system-flow：获取体系级分流结果。",
                ],
                "accept": [
                    "按体系筛选后，问题清单和分流卡片统计一致。",
                    "关闭一个问题后，受影响体系的可通过状态即时更新。",
                    "多体系问题会同时阻断对应体系。",
                    "所有问题处理动作可在历史中追溯。",
                ],
            },
            {
                "name": "认证决定通过/退回",
                "role": "认证决定协调员。",
                "goal": "形成最终认证决定，支持全部通过、全部退回或按体系分流处理。",
                "entry": "顶部任务决策条、右侧“复评判断/认证决定判断”卡片或体系分流卡片。",
                "data": "体系分流结果、专家覆盖、有效问题、任务状态、补正历史。",
                "elements": [
                    "全通过按钮在所有体系满足条件时启用。",
                    "体系级通过按钮在单个体系满足条件时启用。",
                    "退回按钮在存在需补正体系时启用，并要求填写退回原因。",
                    "确认弹窗展示通过体系、退回体系、下一状态和待办对象。",
                ],
                "flow": [
                    "协调员确认专家覆盖和问题清理情况。",
                    "若全部体系可通过，执行认证决定通过。",
                    "若部分体系可通过，执行体系级分流，通过体系记录结论，退回体系生成补正要求。",
                    "若全部不满足，执行认证决定退回。",
                    "系统更新任务或体系状态，并生成流程日志、通知项目管理人员。",
                ],
                "rules": [
                    "全部体系通过后，任务进入已完结。",
                    "任一体系退回时，任务进入二次评定退回，并记录退回体系和原因。",
                    "项目管理人员补正后重新提交直接回到认证决定。",
                    "认证决定操作必须校验当前任务版本，防止并发重复处理。",
                ],
                "api": [
                    "POST /decision-review/{taskId}/pass：全部通过。",
                    "POST /decision-review/{taskId}/systems/{system}/pass：单体系通过。",
                    "POST /decision-review/{taskId}/return：退回一个或多个体系。",
                    "请求需包含 passSystems、returnSystems、returnReason、issueIds、taskVersion。",
                ],
                "accept": [
                    "全部通过后任务状态为已完结，页面只读。",
                    "部分退回后项目管理人员能看到具体退回体系、材料和原因。",
                    "二评退回后的重新提交不会进入初评。",
                    "按钮禁用原因与右侧控制台阻断项一致。",
                ],
            },
        ],
    )


def add_appendices(doc):
    doc.add_page_break()
    add_heading(doc, "8. 接口清单建议", 1)
    add_table(doc, ["接口类别", "接口示例", "说明"], [
        ("任务", "GET /archive-tasks/{taskId}", "获取任务基础信息、状态、当前节点、项目体系和客户信息。"),
        ("材料", "GET /archive-tasks/{taskId}/materials", "获取 MSF11-20 材料项、文件、覆盖体系和缺失状态。"),
        ("完整性", "GET /archive-tasks/{taskId}/completion", "计算各体系必交材料完整性和阻断原因。"),
        ("文件", "POST /archive-tasks/{taskId}/materials/{materialId}/files", "上传或关联资料库文件。"),
        ("资料库", "GET /library/files", "按客户、项目、体系、文件类型和环节查询资料库文件。"),
        ("初评", "POST /first-review/{taskId}/pass", "初评通过。"),
        ("初评", "POST /first-review/{taskId}/return", "初评退回项目管理人员补正。"),
        ("专家评定", "POST /expert-review/{assignmentId}/submit", "专家提交个人评定。"),
        ("问题", "POST /review-issues", "新增材料项级别问题。"),
        ("认证决定", "POST /decision-review/{taskId}/pass", "认证决定全部通过。"),
        ("认证决定", "POST /decision-review/{taskId}/return", "认证决定退回一个或多个体系。"),
        ("日志", "GET /archive-tasks/{taskId}/logs", "查看流程日志、问题历史和关键操作日志。"),
    ], [1500, 3300, 4560])

    add_heading(doc, "9. 验收重点", 1)
    checks = [
        "所有页面默认体系为 Q/E/S/F，且体系名称、颜色和标签规则一致。",
        "材料清单总数为 48 项，分组和排序以 MSF11-20 为准。",
        "评定状态和问题均落在材料清单项级别，不出现在具体文件行评定列。",
        "归档提交、初评、专家评定、认证决定四类角色的权限边界清晰且不可越权。",
        "初评退回和认证决定退回的重新提交路径不同，状态机验证通过。",
        "认证决定支持体系级分流，部分体系通过和部分体系退回的结果可追溯。",
        "所有删除、退回、撤回、通过和提交动作都有自定义确认、日志和必要原因。",
        "按钮禁用均有可读原因，且原因与后端校验结果一致。",
    ]
    for item in checks:
        add_bullet(doc, item)

    add_heading(doc, "10. 待确认问题", 1)
    add_table(doc, ["编号", "问题", "建议默认口径"], [
        ("Q1", "专家提交前是否强制要求所有授权材料均已查看？", "建议不强制，但提交确认弹窗提示未查看材料；可配置为强制。"),
        ("Q2", "认证决定阶段是否使用“二次评定”还是“认证决定”作为系统菜单名称？", "建议菜单和文案统一为“认证决定”，代码状态可兼容原二次评定命名。"),
        ("Q3", "二评退回后是否只允许补正退回体系相关材料？", "建议默认优先定位退回材料，同时保留补充其他材料能力，具体由权限配置控制。"),
        ("Q4", "专家替换后原专家问题是否保留？", "建议保留并标记来源专家，后续由协调员处理或转交新专家复核。"),
        ("Q5", "材料完整性是否只校验必交材料，还是适用时材料也需要人工确认？", "建议提交硬校验只校验必交材料；适用时材料以风险提示方式展示。"),
    ], [700, 4860, 3800])


def main():
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)
    add_description_template(doc)
    add_modules(doc)
    add_appendices(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
