from pathlib import Path
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("7-16状态机说明（Word版）.docx")
FONT_PATH = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "5F6B7A"
GRID = "CBD5E1"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GREEN_FILL = "E6F7F1"
AMBER_FILL = "FFFBE6"
RED_FILL = "FFF1F2"


def set_run_font(run, size=None, bold=None, color=None, east_asia="Arial Unicode MS"):
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)


def set_paragraph_border_bottom(paragraph, color=BLUE, size=12, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, label, text, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=GRID, size=5)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label + "：")
    set_run_font(r1, size=10.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=HEADER_FILL, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_margins(cell)
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_flow(doc, text, fill=LIGHT_FILL):
    add_callout(doc, "流程", text, fill=fill)


def embed_font(docx_path, font_path, font_name="Arial Unicode MS"):
    """Embed an obfuscated TrueType font so Chinese renders consistently."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ET.register_namespace("w", w_ns)
    ET.register_namespace("r", r_ns)

    with tempfile.TemporaryDirectory(prefix="docx_font_embed_") as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as src:
            src.extractall(tmp_path)

        font_table_path = tmp_path / "word" / "fontTable.xml"
        rels_path = tmp_path / "word" / "_rels" / "fontTable.xml.rels"
        rels_path.parent.mkdir(parents=True, exist_ok=True)
        fonts_dir = tmp_path / "word" / "fonts"
        fonts_dir.mkdir(parents=True, exist_ok=True)

        font_key = uuid.uuid4()
        key_bytes = font_key.bytes[::-1]
        raw = bytearray(font_path.read_bytes())
        for i in range(min(32, len(raw))):
            raw[i] ^= key_bytes[i % 16]
        (fonts_dir / "font1.odttf").write_bytes(raw)

        font_tree = ET.parse(font_table_path)
        font_root = font_tree.getroot()
        font_node = None
        for node in font_root.findall(f"{{{w_ns}}}font"):
            if node.get(f"{{{w_ns}}}name") == font_name:
                font_node = node
                break
        if font_node is None:
            font_node = ET.SubElement(font_root, f"{{{w_ns}}}font")
            font_node.set(f"{{{w_ns}}}name", font_name)
        for old in font_node.findall(f"{{{w_ns}}}embedRegular"):
            font_node.remove(old)
        embed = ET.SubElement(font_node, f"{{{w_ns}}}embedRegular")
        embed.set(f"{{{r_ns}}}id", "rIdFont1")
        embed.set(f"{{{w_ns}}}fontKey", "{" + str(font_key).upper() + "}")
        font_tree.write(font_table_path, encoding="UTF-8", xml_declaration=True)

        if rels_path.exists():
            rel_tree = ET.parse(rels_path)
            rel_root = rel_tree.getroot()
        else:
            rel_root = ET.Element(f"{{{rel_ns}}}Relationships")
            rel_tree = ET.ElementTree(rel_root)
        for old in list(rel_root):
            if old.get("Id") == "rIdFont1":
                rel_root.remove(old)
        rel = ET.SubElement(rel_root, f"{{{rel_ns}}}Relationship")
        rel.set("Id", "rIdFont1")
        rel.set("Type", r_ns + "/font")
        rel.set("Target", "fonts/font1.odttf")
        ET.register_namespace("", rel_ns)
        rel_tree.write(rels_path, encoding="UTF-8", xml_declaration=True)

        settings_path = tmp_path / "word" / "settings.xml"
        settings_tree = ET.parse(settings_path)
        settings_root = settings_tree.getroot()
        if settings_root.find(f"{{{w_ns}}}embedTrueTypeFonts") is None:
            settings_root.insert(0, ET.Element(f"{{{w_ns}}}embedTrueTypeFonts"))
        settings_tree.write(settings_path, encoding="UTF-8", xml_declaration=True)

        content_types_path = tmp_path / "[Content_Types].xml"
        ct_tree = ET.parse(content_types_path)
        ct_root = ct_tree.getroot()
        has_font_default = any(
            node.get("Extension") == "odttf" for node in ct_root.findall(f"{{{ct_ns}}}Default")
        )
        if not has_font_default:
            node = ET.SubElement(ct_root, f"{{{ct_ns}}}Default")
            node.set("Extension", "odttf")
            node.set("ContentType", "application/vnd.openxmlformats-officedocument.obfuscatedFont")
        ET.register_namespace("", ct_ns)
        ct_tree.write(content_types_path, encoding="UTF-8", xml_declaration=True)

        temp_docx = tmp_path / "embedded.docx"
        with zipfile.ZipFile(temp_docx, "w", zipfile.ZIP_DEFLATED) as dst:
            for part in tmp_path.rglob("*"):
                if part.is_file() and part != temp_docx:
                    dst.write(part, part.relative_to(tmp_path).as_posix())
        docx_path.write_bytes(temp_docx.read_bytes())


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial Unicode MS"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Arial Unicode MS"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = header.add_run("体系认证项目｜业务状态机需求说明")
set_run_font(hr, size=9, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("业务状态机需求说明")
set_run_font(r, size=23, bold=True, color=INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("体系归档、初次评定、专家评定、认证决定及问题整改闭环")
set_run_font(r, size=13, color=MUTED)

metadata = [
    ("来源", "7-16状态机图.html"),
    ("适用范围", "体系认证归档与评定业务"),
    ("业务口径", "认证决定（二次评定）通过后，本阶段业务结束并进入已完结"),
    ("文档用途", "可直接复制到产品需求文档，表格与文字均可编辑"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(label + "：")
    set_run_font(r1, size=10.5, bold=True, color=INK)
    r2 = p.add_run(value)
    set_run_font(r2, size=10.5, color=INK)
rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(7)
rule.paragraph_format.space_after = Pt(8)
set_paragraph_border_bottom(rule)

add_callout(doc, "核心原则", "状态必须按业务对象分别管理，不将归档任务、初次评定、专家任务、认证决定、具体问题和整体项目的状态混入同一字段。问题状态独立流转，任务是否可通过由问题闭环结果聚合判断。")

add_heading(doc, "1. 状态机范围与对象", 1)
scope_rows = [
    ("归档材料提交任务", "项目管理人员", "准备材料、提交初评；二评退回后补充材料并直接回到认证决定"),
    ("初次评定任务", "初次评定管理员/协调员", "接收任务、指派专家、处理问题并作出初评结论"),
    ("专家评定任务", "专家", "开始评审、登记问题、完成评审"),
    ("认证决定（二次评定）任务", "二次评定管理员/协调员", "接收任务、指派专家、按体系放行或退回"),
    ("具体问题", "专家、协调员、项目管理人员", "登记、退回、整改、复核、通过或删除"),
    ("整体项目", "系统聚合", "根据各阶段完成情况进入下一阶段或已完结"),
]
add_table(doc, ["业务对象", "主要角色", "状态机职责"], scope_rows, [2250, 2250, 4860], font_size=9.5)

add_heading(doc, "2. 状态数据模型", 1)
add_body(doc, "每个状态应至少具备以下配置字段。建议后端使用固定编码，页面仅展示状态名称；可执行操作与角色权限应由状态机配置和权限矩阵共同约束。")
field_rows = [
    ("状态编码", "后端、接口和日志使用的固定编码，不随页面文案变化"),
    ("状态名称", "页面展示文案"),
    ("所属对象", "归档任务、初评任务、专家任务、认证决定任务、具体问题、体系或整体项目"),
    ("进入条件", "进入本状态必须满足的前置条件"),
    ("可执行操作", "当前状态允许展示并执行的业务按钮"),
    ("操作角色", "允许执行操作的角色或岗位"),
    ("目标状态", "操作成功后的状态；失败时保持原状态"),
    ("是否生成待办", "是否生成待办、待办接收角色及关闭条件"),
    ("是否允许撤回", "撤回条件、撤回目标状态及是否保留操作历史"),
    ("是否只读", "页面是否锁定，以及允许补充说明或上传附件的范围"),
]
add_table(doc, ["字段", "说明"], field_rows, [2700, 6660], font_size=10)

add_heading(doc, "3. 归档材料提交任务状态机", 1)
add_flow(doc, "草稿 →（提交初评）→ 初次评定中 →（初评问题全部通过）→ 认证决定中 →（认证决定问题全部通过）→ 已完结", fill=GREEN_FILL)
archive_rows = [
    ("草稿", "提交初评", "项目管理人员", "初次评定中", "Q/E/S/F 各体系必备材料齐全"),
    ("初次评定中", "问题未全部通过", "系统聚合", "初次评定中", "保持当前阶段，等待问题整改闭环"),
    ("初次评定中", "初评通过", "初次评定管理员", "认证决定中", "初评有效问题全部通过"),
    ("认证决定中", "问题未全部通过", "系统聚合", "认证决定中", "保持当前阶段，等待问题整改闭环"),
    ("认证决定中", "认证决定通过", "二次评定管理员", "已完结", "可放行体系完成决定；本阶段业务结束"),
    ("认证决定退回", "补充材料并再次提交", "项目管理人员", "认证决定中", "不再经过初次评定"),
]
add_table(doc, ["当前状态", "操作/条件", "执行角色", "目标状态", "业务规则"], archive_rows, [1500, 1900, 1500, 1500, 2960], font_size=8.8)
add_callout(doc, "口径说明", "源图中还绘制了“是否涉及证书、证书签发/不予签发、待推送/已推送”等下游节点。根据当前项目统一口径，这些节点不作为本阶段主任务状态；如后续纳入证书或推送子流程，应建立独立对象状态机。", fill=AMBER_FILL)

add_heading(doc, "4. 初次评定任务状态机", 1)
add_flow(doc, "待接收 →（接收）→ 待指派专家 →（完成指派）→ 初次评定中 →（问题全部通过）→ 初评通过")
initial_rows = [
    ("待接收", "接收", "初次评定管理员", "待指派专家", "接收后生成指派专家待办"),
    ("待指派专家", "完成指派", "初次评定管理员", "初次评定中", "专家覆盖所评体系，且至少有一名有效专家"),
    ("初次评定中", "专家提出问题", "专家", "初次评定中", "问题独立进入问题状态机"),
    ("初次评定中", "问题未全部通过", "系统聚合", "初次评定中", "通过按钮保持禁用并展示阻断原因"),
    ("初次评定中", "通过初评", "初次评定管理员", "初评通过", "所有专家任务完成且有效问题全部通过"),
    ("初评通过", "二评退回需初评确认", "初次评定管理员", "二评整改待初评确认/初评通过", "仅在二评回退链路需要确认时启用"),
]
add_table(doc, ["当前状态", "操作/事件", "执行角色", "目标状态", "进入/放行条件"], initial_rows, [1500, 1800, 1650, 1550, 2860], font_size=8.8)

add_heading(doc, "5. 专家评定任务状态机", 1)
add_flow(doc, "待评审 →（开始评审）→ 初次评审中 →（问题全部整改完成）→ 已完成")
expert_task_rows = [
    ("待评审", "开始评审", "专家", "初次评审中", "专家已被指派且任务有效"),
    ("初次评审中", "登记并提交问题", "专家", "初次评审中", "问题进入具体问题状态机"),
    ("初次评审中", "问题全部整改完成", "系统聚合/专家", "已完成", "本专家提出的有效问题全部通过或已删除"),
]
add_table(doc, ["当前状态", "操作/事件", "执行角色", "目标状态", "进入/完成条件"], expert_task_rows, [1650, 1900, 1650, 1550, 2610], font_size=9)
add_callout(doc, "注意", "“评审中”包含单个专家任务的问题登记、整改和复核过程；只有该专家任务下的有效问题全部闭环后，专家任务才可置为“已完成”。")

add_heading(doc, "6. 认证决定（二次评定）任务状态机", 1)
add_flow(doc, "待接收 →（接收）→ 待指派专家 →（完成指派）→ 认证决定中 →（全部可放行）→ 认证决定通过 → 已完结", fill=GREEN_FILL)
decision_rows = [
    ("待接收", "接收", "二次评定管理员", "待指派专家", "接收后生成指派专家待办"),
    ("待指派专家", "完成指派", "二次评定管理员", "认证决定中", "专家覆盖 Q/E/S/F 中本任务涉及的体系"),
    ("认证决定中", "专家提出问题", "专家", "认证决定中", "问题按体系归属进入问题状态机"),
    ("认证决定中", "问题未全部通过", "系统聚合", "认证决定中", "对应体系不可放行"),
    ("认证决定中", "全部可放行", "二次评定管理员", "认证决定通过", "体系级问题已清理且满足放行条件"),
    ("认证决定通过", "结束业务", "系统", "已完结", "二次评定通过后本阶段结束"),
]
add_table(doc, ["当前状态", "操作/事件", "执行角色", "目标状态", "进入/放行条件"], decision_rows, [1500, 1800, 1650, 1550, 2860], font_size=8.8)
add_callout(doc, "体系级分流", "二次评定应按 Q/E/S/F 体系分别计算“可通过体系”和“需退回体系”。具体问题挂在材料项/体系上，评定结论落在体系或任务层级，不落在具体文件行。", fill=AMBER_FILL)

add_heading(doc, "7. 具体问题状态机", 1)
add_body(doc, "问题状态独立于任务状态。任务是否可通过，只聚合“有效且未删除”的问题；问题删除、无需解决和复核退回均必须保留历史记录。")

add_heading(doc, "7.1 项目管理人员/协调员整改视角", 2)
add_flow(doc, "待退回/问题退回 → 待整改 →（整改完成并提交）→ 待复核 →（复核）→ 复核通过；复核不通过则进入复核退回并重新整改", fill=RED_FILL)
rectify_rows = [
    ("待退回", "问题退回", "协调员", "待整改", "填写退回意见并通知项目管理人员"),
    ("待整改", "整改完成并提交", "项目管理人员", "待复核", "补充整改说明及必要材料"),
    ("待整改", "认为无需解决", "项目管理人员/协调员", "无需解决", "必须填写理由，仍需复核"),
    ("待复核/无需解决", "复核", "问题提出人或授权复核人", "是否通过", "只允许处理本人或授权范围内的问题"),
    ("是否通过", "通过", "复核人", "复核通过", "问题不再阻断任务通过"),
    ("是否通过", "复核不通过，退回", "复核人", "复核退回", "填写复核意见并重新生成整改待办"),
    ("复核退回", "重新整改", "项目管理人员", "待整改", "保留上一轮整改和复核记录"),
]
add_table(doc, ["当前状态", "操作", "角色", "目标状态", "约束/结果"], rectify_rows, [1550, 1850, 1650, 1450, 2860], font_size=8.8)

add_heading(doc, "7.2 专家登记与复核视角", 2)
add_flow(doc, "草稿 →（提交协调员）→ 已提交协调员 → 待整改 → 待复核 → 复核通过；草稿或允许撤回的记录可删除", fill=LIGHT_FILL)
expert_issue_rows = [
    ("草稿", "登记问题", "专家", "草稿", "问题尚未生效，不阻断任务"),
    ("草稿", "提交协调员", "专家", "已提交协调员", "校验问题类型、所属材料项、适用体系、问题描述"),
    ("草稿/已撤回", "删除", "专家", "已删除", "软删除；保留审计历史"),
    ("已提交协调员", "退回整改", "协调员", "待整改", "生成项目管理人员整改待办"),
    ("已提交协调员", "退回专家修改", "协调员", "草稿", "说明退回原因，专家可编辑后再次提交"),
    ("待整改", "整改并提交", "项目管理人员", "待复核", "记录整改说明、附件和提交时间"),
    ("待复核", "复核通过", "专家", "复核通过", "问题闭环"),
    ("待复核", "复核不通过", "专家", "复核退回", "重新进入整改链路"),
]
add_table(doc, ["当前状态", "操作", "角色", "目标状态", "约束/结果"], expert_issue_rows, [1550, 1850, 1450, 1450, 3060], font_size=8.8)

add_heading(doc, "8. 角色权限矩阵（建议基线）", 1)
add_body(doc, "权限由“角色 × 业务对象 × 当前状态”共同决定。下表用于需求评审，最终以组织权限和任务指派关系为准。")
permission_rows = [
    ("项目管理人员", "√", "补充/上传", "—", "—", "—", "提交整改", "—", "—"),
    ("初次评定管理员/协调员", "√", "受控", "√", "√", "√", "退回整改", "授权范围", "通过初评/退回"),
    ("二次评定管理员/协调员", "√", "受控", "√", "√", "√", "退回整改", "授权范围", "放行/退回体系"),
    ("专家", "√", "—", "—", "√", "√", "—", "本人问题", "—"),
    ("系统", "按规则", "—", "—", "—", "—", "生成待办", "聚合结果", "自动流转"),
]
add_table(doc, ["角色", "查看任务", "编辑材料", "指派专家", "登记问题", "整改问题", "退回", "复核", "通过/放行"], permission_rows, [1500, 880, 980, 920, 920, 920, 800, 980, 1460], font_size=8.0)

add_heading(doc, "9. 通用流转与交互规则", 1)
rules = [
    ("R-01", "状态变更必须校验当前状态，避免重复提交或越级流转。"),
    ("R-02", "每次流转记录操作人、操作时间、来源页面、原状态、目标状态、操作意见和附件。"),
    ("R-03", "状态变更、业务数据更新、待办生成/关闭和通知写入应保持事务一致性。"),
    ("R-04", "主按钮不可用时必须展示明确阻断原因，例如“仍有 3 个有效问题待复核”。"),
    ("R-05", "撤回仅在下游角色尚未处理且未产生不可逆结果时开放；撤回后保留历史记录。"),
    ("R-06", "删除问题采用软删除；已进入整改或复核链路的问题原则上不允许直接物理删除。"),
    ("R-07", "具体文件不是评定结论对象；问题挂在材料清单项/体系上，文件仅作为证据和附件。"),
    ("R-08", "任务汇总时仅统计有效问题；草稿、已删除问题不阻断，待整改、待复核、复核退回问题阻断。"),
    ("R-09", "二次评定退回后，项目管理人员补充材料再次提交，直接回到认证决定，不再经过初次评定。"),
    ("R-10", "模拟体系统一使用 Q/E/S/F：质量、环境、职业健康安全、食品安全；不得使用其他默认体系。"),
]
add_table(doc, ["编号", "规则"], rules, [1200, 8160], font_size=9.5)

add_heading(doc, "10. 验收检查清单", 1)
acceptance_rows = [
    ("对象隔离", "不同任务和问题使用独立状态字段，无混用或覆盖"),
    ("主链路", "归档 → 初评 → 认证决定 → 已完结可完整走通"),
    ("问题闭环", "登记、退回、整改、复核、通过/再次退回均可追溯"),
    ("权限", "无权限角色不可见或不可执行对应操作，接口同步校验"),
    ("阻断", "存在有效未闭环问题时，任务通过/体系放行被阻断且原因明确"),
    ("二评回退", "补充材料再次提交后直接返回认证决定"),
    ("体系口径", "仅使用 Q/E/S/F，不使用其他默认体系"),
    ("历史与待办", "状态历史、待办生成/关闭、通知结果一致"),
]
add_table(doc, ["验收项", "通过标准"], acceptance_rows, [2400, 6960], font_size=9.5)

doc.core_properties.title = "体系归档业务状态机说明"
doc.core_properties.subject = "归档、初次评定、专家评定、认证决定及问题整改状态机"
doc.core_properties.author = "Codex"
doc.core_properties.keywords = "体系认证, 状态机, 初次评定, 二次评定, 问题整改"
doc.save(OUT)
if FONT_PATH.exists():
    embed_font(OUT, FONT_PATH)
print(OUT.resolve())
