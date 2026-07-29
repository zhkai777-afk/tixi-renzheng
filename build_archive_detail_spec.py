from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("/Users/zhazhakai777/Documents/体系认证/归档材料提交详情页功能描述.docx")
W = 9360
BLUE, DARK, INK, MUTED = "2E74B5", "1F4D78", "202124", "5F6368"
LIGHT, LIGHT_BLUE, BORDER = "F2F4F7", "E8EEF5", "D9E1EA"
FONT = "Arial"


def font(run, size=10.5, color=INK, bold=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fs = rpr.get_or_add_rFonts()
    fs.set(qn("w:ascii"), FONT)
    fs.set(qn("w:hAnsi"), FONT)


def style_font(style, size, color=INK, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fs = style.element.get_or_add_rPr().get_or_add_rFonts()
    fs.set(qn("w:ascii"), FONT)
    fs.set(qn("w:hAnsi"), FONT)


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    n = pr.find(qn("w:shd"))
    if n is None:
        n = OxmlElement("w:shd")
        pr.append(n)
    n.set(qn("w:fill"), fill)


def no_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    n = OxmlElement("w:tblHeader")
    n.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(n)


def geometry(table, widths):
    assert sum(widths) == W
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tw.getparent() is None:
        pr.append(tw)
    tw.set(qn("w:w"), str(W)); tw.set(qn("w:type"), "dxa")
    ti = pr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if ti.getparent() is None:
        pr.append(ti)
    ti.set(qn("w:w"), "120"); ti.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for c in list(grid):
        grid.remove(c)
    for x in widths:
        c = OxmlElement("w:gridCol"); c.set(qn("w:w"), str(x)); grid.append(c)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell._tc.get_or_add_tcPr()
            cw = cp.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if cw.getparent() is None:
                cp.append(cw)
            cw.set(qn("w:w"), str(widths[i])); cw.set(qn("w:type"), "dxa")
            mar = cp.find(qn("w:tcMar")) or OxmlElement("w:tcMar")
            if mar.getparent() is None:
                cp.append(mar)
            for name, val in (("top",80),("bottom",80),("start",120),("end",120)):
                n = mar.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
                if n.getparent() is None:
                    mar.append(n)
                n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, fill=LIGHT):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], fill)
        p = t.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), 9.5, DARK, True)
    for vals in rows:
        cells = t.add_row().cells
        no_split(t.rows[-1])
        for i, v in enumerate(vals):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font(p.add_run(str(v)), 9.2)
    geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    font(p.add_run(text), 10.5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    font(p.add_run(text), 10.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def callout(doc, label, text, color=BLUE, fill="F4F6F9"):
    t = doc.add_table(rows=1, cols=1)
    no_split(t.rows[0]); geometry(t, [W])
    c = t.cell(0, 0); shade(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "："), 10.5, color, True)
    font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def op(doc, name, vals):
    heading(doc, name, 3)
    labels = ["触发入口","显示/启用条件","前置条件","校验规则","成功状态","失败提示","确认与日志"]
    for label, value in zip(labels, vals):
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.left_indent = Inches(.18)
        p.paragraph_format.first_line_indent = Inches(-.18)
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(label + "："), 10, DARK, True)
        font(p.add_run(value), 10)


def page_number(p):
    r = p.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = "PAGE"
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "end")
    r._r.extend([a,b,c]); font(r, 8.5, MUTED)


def setup():
    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.right_margin = s.bottom_margin = s.left_margin = Inches(1)
    s.header_distance = s.footer_distance = Inches(.492)
    style_font(d.styles["Normal"], 10.5)
    d.styles["Normal"].paragraph_format.space_after = Pt(6)
    d.styles["Normal"].paragraph_format.line_spacing = 1.1
    style_font(d.styles["Body Text"], 10.5)
    d.styles["Body Text"].paragraph_format.space_after = Pt(6)
    d.styles["Body Text"].paragraph_format.line_spacing = 1.1
    for n, size, color, before, after in ((1,16,BLUE,16,8),(2,13,BLUE,12,6),(3,11.5,DARK,8,4)):
        st = d.styles[f"Heading {n}"]; style_font(st,size,color,True)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    style_font(d.styles["List Bullet"], 10.3)
    d.styles["List Bullet"].paragraph_format.left_indent = Inches(.38)
    d.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-.18)
    hp = s.header.paragraphs[0]; font(hp.add_run("体系认证平台  /  功能说明书"),8.5,MUTED,True)
    fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("归档材料提交详情页  ·  第 "),8.5,MUTED); page_number(fp); font(fp.add_run(" 页"),8.5,MUTED)
    return d


def build():
    d = setup()
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); font(p.add_run("功能说明书"),11,BLUE,True)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(5); font(p.add_run("归档材料提交详情页"),24,INK,True)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(14); font(p.add_run("归档材料补充、完整性检查与初次评定提交"),13,MUTED)
    for k,v in [
        ("文档对象","归档材料提交-详情-10体系重构版.html"),
        ("页面角色","项目管理人员"),
        ("业务范围","查看带入材料、补充文件、检查 Q/E/S/F 完整性、提交初次评定"),
        ("数据口径","MSF11-20 审核材料归档/上报清单，共 48 项"),
        ("版本日期","V1.0 / 2026-07-23"),
    ]:
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(k+"："),10.5,INK,True); font(p.add_run(v),10.5)
    callout(d,"页面定位","项目管理人员在此核对前序环节自动带入的归档证据，对缺失材料执行本地上传或资料库关联；所有涉及体系适用必交材料达到 100% 后，提交进入初次评定。")

    heading(d,"1. 功能基本信息",1)
    add_table(d,["项目","说明"],[
        ("功能名称","归档材料提交详情页"),
        ("使用角色","项目管理人员；已完结或无编辑权限时以只读模式查看。"),
        ("业务目标","补齐归档材料、确认体系覆盖与提交条件，将完整材料提交初次评定。"),
        ("入口","材料归档管理列表 → 草稿/退回任务详情；也可由项目待办进入。"),
        ("前置状态","已生成归档任务；系统已聚合前序环节附件和在线表单；用户具备任务权限。"),
        ("数据权限","按任务、机构、部门和项目授权过滤；文件预览、删除、上传均服务端鉴权。"),
        ("材料口径","MSF11-20 共 48 项：通用材料 32、一阶段材料 7、二阶段及其他审核材料 9。"),
    ],[1900,7460],LIGHT_BLUE)

    heading(d,"2. 页面展示",1)
    heading(d,"2.1 页面区域与信息优先级",2)
    add_table(d,["区域","展示内容","业务作用"],[
        ("顶部任务信息","返回、标题、TASK-2026-001、客户A、Q/E/S/F、当前视图、提示与摘要卡。","先确认任务范围、当前状态和整体完整度。"),
        ("顶部操作区","操作指引、模拟补齐（仅原型）、保存草稿、提交初次评定。","承载页面级动作；正式系统不保留模拟按钮。"),
        ("主工作区标题","材料清单视图/全部文件视图、全部/缺失/已带入、体系筛选、搜索。","快速缩小材料范围并切换阅读口径。"),
        ("材料台账","48 项材料：名称、适用体系、文件数量、覆盖状态、补充操作。","以材料清单项为完整性判断对象。"),
        ("文件展开区","具体文件名称、所属材料、类型、来源、上传时间、预览/删除。","查看材料项下证据及来源。"),
        ("上传补充面板","待补充文件、目标材料、覆盖体系、取消/确认补充。","把本地文件准确关联到材料项和体系。"),
        ("右侧提交控制台","任务状态、Q/E/S/F 各体系完成度、缺失项与提交条件。","解释是否可提交以及阻断原因。"),
        ("弹窗与反馈","资料库选择、删除确认、提交确认、自定义 toast、操作指引。","提供选择、确认、反馈和防误操作。"),
    ],[1650,4100,3610])
    heading(d,"2.2 视图与筛选规则",2)
    add_table(d,["功能","规则"],[
        ("材料清单视图","展示当前范围内全部材料项；可展开文件、上传文件或从资料库选择。"),
        ("全部文件视图","仅展示已有具体文件，不生成缺失材料行；显示文件数及来源材料数。"),
        ("快速筛选-全部","材料台账展示全部材料；文件视图只展示已有文件。"),
        ("快速筛选-缺失","只展示缺失或部分覆盖材料；文件视图显示引导空状态。"),
        ("快速筛选-已带入","只展示已有文件且覆盖状态为已齐全的材料。"),
        ("体系范围","支持 Q/E/S/F 多选；材料项任一适用体系命中即展示，文件按覆盖体系过滤。"),
        ("搜索","对材料名称、文件名、来源、分类等实时模糊匹配；去除首尾空格。"),
        ("材料分组","按通用材料、一阶段材料、二阶段及其他审核材料分组并支持展开/收起。"),
    ],[2000,7360])
    heading(d,"2.3 材料台账字段",2)
    add_table(d,["字段","展示规则"],[
        ("序号","按 MSF11-20 sortOrder 排序；筛选后可显示当前结果序号。"),
        ("材料清单项","使用权威材料名称；第一行为名称，不与体系徽标挤在同一行。"),
        ("适用体系","Q/E/S/F 弱化色徽标；非当前筛选体系可灰化。"),
        ("文件数量","显示 N 个文件或未上传；多个文件不增加材料项完成计数。"),
        ("状态","缺失、部分覆盖、已齐全；按材料适用体系与文件覆盖体系实时计算。"),
        ("操作","上传文件、资料库选择、查看/收起文件；只读状态仅保留查看。"),
    ],[2000,7360])
    heading(d,"2.4 具体文件字段",2)
    add_table(d,["字段","展示位置/规则"],[
        ("具体文件名称","主信息，14px/加粗；支持预览。"),
        ("所属材料","作为文件名下方辅助信息，不单列。"),
        ("文件类型/来源/上传时间","作为辅助信息展示，便于追溯。"),
        ("覆盖体系","应使用 Q/E/S/F 徽标；本原型部分视图未单列，正式版本应保持可识别。"),
        ("操作","预览、删除；删除必须二次确认并重新计算完整性。"),
    ],[2200,7160])
    callout(d,"重要口径","评定状态和问题落在 MSF11-20 材料清单项，不在具体文件行设置“评定状态”列；具体文件仅是材料项下的证据。","C46000","FFF8E7")

    heading(d,"3. 数据来源",1)
    body(d,"页面以归档任务、MSF11-20 材料目录、前序业务文件、资料库文件、手工上传记录及流程状态为数据基础。当前原型使用静态模拟文件和前端内存状态；正式系统应由服务端聚合、持久化并进行版本控制。")
    add_table(d,["字段/数据","来源对象","来源字段/算法","实时","空值/失败处理","权限"],[
        ("任务信息","归档任务/认证项目","task_no、customer、systems、audit_type、status","否","任务缺失则返回列表并提示","任务授权"),
        ("48项材料","materials-catalog.json","id、category、materialName、applicableSystems、requirements、sortOrder","否","目录加载失败阻断编辑/提交","系统公共配置"),
        ("前序带入文件","业务附件/在线表单","material_id、file_id、source_step、systems、valid_flag","是","失效引用显示异常并重算","项目授权"),
        ("本地上传","归档文件服务","file_id、material_id、systems、uploader、time、scan_status","是","上传/扫描失败不计完成度","编辑权限"),
        ("资料库文件","客户资料库","library_file_id、file_type、systems、step、uploader","是","无结果显示空状态","资料库+项目权限"),
        ("材料覆盖状态","完整性服务","适用体系 - 有效文件覆盖体系","是","计算失败禁用提交","任务授权"),
        ("体系完整度","完整性服务","各体系已完成适用必交项/适用必交项","是","显示计算中/失败及重试","任务授权"),
        ("流程状态","流程实例","draft、initial_review、returned、completed 等","是","未知状态只读并告警","流程权限"),
    ],[1200,1400,2500,650,2150,1460],LIGHT_BLUE)
    heading(d,"3.1 完整性计算规则",2)
    for x in [
        "材料目录必须使用 MSF11-20 48 项，不允许手写其他总数；适用性由认证类型、Q/E/S/F 项目范围及 requirements 共同决定。",
        "材料项存在有效文件且其覆盖体系包含该材料在当前项目中的全部适用体系时，状态为已齐全；覆盖部分体系时为部分覆盖；无有效文件时为缺失。",
        "提交门槛按各体系适用必交材料 100% 计算，不能仅以“48 项都有文件”或文件总数作为依据。",
        "删除、扫描失败、资料库引用失效、评定退回均须触发重新计算；计算完成前禁止提交。",
    ]: bullet(d,x)
    callout(d,"原型差异","当前 HTML 将“提交初次评定”始终设为可用，并在缺失时提示仍可提交；该行为与项目规则冲突。正式实现必须在所有涉及体系适用必交材料齐全后启用提交，并在禁用时显示具体缺失原因。","B42318","FFF1F0")

    heading(d,"4. 操作规则",1)
    operations = [
        ("切换视图",["点击材料清单视图/全部文件视图","用户可查看任务","任务已加载","保留体系、状态、搜索条件；清理未确认上传预览","刷新工作区并更新当前视图名称","加载失败保持原视图","无需确认；记录页面行为"]),
        ("筛选与搜索",["点击全部/缺失/已带入、勾选体系或输入搜索词","任务已加载","无","多个条件按 AND 组合；体系内部按 OR 命中","材料树、台账、文件统计和右侧控制台同步更新","无结果显示业务空状态","无需确认"]),
        ("展开材料文件",["点击材料行或查看文件","材料项可见","无","仅切换当前材料展开状态；文件按当前体系范围过滤","显示该材料下具体文件及数量","无文件时给出上传/资料库入口","无需确认"]),
        ("本地上传",["材料行点击上传文件，选择一个或多个本地文件","草稿/退回整改且有编辑权限","已选具体材料","校验扩展名、大小、数量、病毒扫描；至少选择一个适用覆盖体系；禁止重复文件按产品规则处理","文件入库并关联材料项/体系；重算完整性；展示来源与时间","逐文件提示失败原因；成功文件可保留，事务策略需明确","确认补充；记录操作人、文件哈希、材料项、体系和结果"]),
        ("从资料库选择",["材料行点击资料库选择","有编辑和资料库查看权限","已选具体材料","可按分类、关键字、体系、文件类型、上传环节筛选；至少选择一项；校验文件有效性与体系适用性","建立引用关系并重算完整性","无结果显示空状态；失效/无权限文件拒绝关联","确认选择（N）；记录来源文件和关联关系"]),
        ("预览文件",["点击预览","有文件查看权限","文件存在且状态有效","按文件类型选择预览器；禁止直接暴露存储地址","打开预览并记录访问审计","不支持或文件失败时提供明确提示/受控下载","无需确认"]),
        ("删除文件",["点击删除","草稿/退回整改且有删除权限；非系统锁定文件","文件存在","服务端校验任务状态、文件版本和引用关系","删除/解除关联，重算材料和体系完整性","状态已变化或无权限时拒绝并刷新","二次确认；记录删除人、时间、材料项和文件"]),
        ("保存草稿",["点击保存草稿","任务可编辑","无未完成上传事务","持久化当前材料关联和补充记录；不得仅做前端提示","保存成功且状态仍为草稿/整改","失败时保留界面数据并允许重试","无需二次确认；记录保存结果"]),
        ("提交初次评定",["点击提交初次评定","项目管理人员有权限；状态为草稿/初评退回；按钮仅在条件满足时启用","Q/E/S/F 各涉及体系适用必交材料均为100%；无扫描失败/上传中/统计失败","服务端原子校验权限、任务版本、材料有效性、完整性、重复请求","状态进入初次评定/未接收；页面转只读；创建初评待办","返回具体体系和材料项；不得部分提交","自定义确认弹窗；幂等键；完整审计日志"]),
        ("操作指引",["点击操作指引","页面可用","无","按4步定位顶部、左侧筛选、中间工作区、提交按钮；焦点和遮罩可访问","支持上一步/下一步/结束","目标元素不可见时跳过或提示","无需业务日志"]),
    ]
    for n,v in operations: op(d,n,v)

    heading(d,"4.1 状态与页面模式",2)
    add_table(d,["任务状态","模式","可用动作","说明"],[
        ("草稿","编辑","筛选、上传、资料库、预览、删除、保存、提交","提交条件不足时显示缺失原因。"),
        ("初次评定中/未接收","只读或有限撤回","预览、查看完整性；撤回由列表/流程规则提供","详情页不继续编辑材料。"),
        ("初次评定中/已接收","只读","预览、查看","禁止撤回和材料变更。"),
        ("初次评定退回","整改","上传、资料库、预览、删除、重新提交","展示退回原因和受影响材料项。"),
        ("二次评定退回","整改","上传、资料库、预览、删除、提交整改","提交后进入二评整改待初评确认。"),
        ("二评整改待初评确认","只读/待确认","预览、查看","等待初评管理员确认。"),
        ("已完结","只读","预览、查看来源与历史","禁止任何材料变更。"),
    ],[2000,1600,3100,2660])

    heading(d,"5. 异常场景",1)
    add_table(d,["异常","触发情形","处理规则"],[
        ("材料目录加载失败","materials-catalog.json/目录服务不可用","显示重试，禁用上传和提交；不得使用不完整 fallback 继续正式业务。"),
        ("上传失败","网络中断、格式/大小超限、扫描失败","逐文件提示；失败文件不计完整度；允许重试。"),
        ("重复文件","同一哈希已存在于材料项或任务","提示已存在并由规则决定阻止或复用，不重复占用存储。"),
        ("资料库引用失效","源文件删除、过期或权限收回","标记失效，重新计算完整度并提示重新补充。"),
        ("覆盖体系不适用","用户勾选材料适用范围外体系","前端不展示，服务端再次拒绝。"),
        ("状态已变化","编辑期间任务被提交、接收或完结","基于版本号拒绝变更，刷新为最新只读状态。"),
        ("权限被收回","操作前权限变化","服务端拒绝并隐藏编辑入口，不泄露文件信息。"),
        ("删除后不完整","删除最后一个有效覆盖文件","立即降级状态、显示缺失并禁用提交。"),
        ("统计服务失败","完整性计算超时/异常","显示计算失败，提交禁用，提供重试。"),
        ("重复提交","双击或网络重试","幂等处理，仅创建一次状态流转和待办。"),
        ("并发提交/删除","一个用户提交时另一个删除","以后端原子事务和版本号为准，失败方刷新。"),
        ("预览失败","格式不支持、文件损坏","提示原因；必要时提供受控下载，不影响其他文件。"),
        ("空状态","筛选无材料或无文件","说明缺什么、为何无内容、如何调整或补充。"),
    ],[1700,3000,4660])

    heading(d,"6. 验收标准（Given–When–Then）",1)
    scenarios = [
        ("场景 1：页面初始化","给定：草稿任务 TASK-2026-001，项目范围为 Q/E/S/F。","当：用户进入详情页。","那么：加载 48 项材料及已有文件，显示任务信息、材料汇总和四体系完整度。"),
        ("场景 2：目录口径","给定：权威目录包含 32+7+9 项。","当：页面加载成功。","那么：总计严格为 48 项，名称、分类、排序与 materials-catalog.json 一致。"),
        ("场景 3：组合筛选","给定：存在 E 专属缺失材料。","当：选择 E 和缺失并搜索材料名。","那么：仅显示同时满足条件的材料，台账、材料树和统计同步。"),
        ("场景 4：全部文件视图","给定：部分材料已有文件。","当：切换全部文件视图。","那么：只展示具体文件及所属材料/来源，不生成缺失材料行。"),
        ("场景 5：本地上传成功","给定：E/S 材料仍缺覆盖，用户有编辑权限。","当：上传有效 PDF 并勾选 E/S 后确认。","那么：文件关联该材料与 E/S，状态和体系完整度实时更新并写审计日志。"),
        ("场景 6：上传必填失败","给定：用户已选文件但未选覆盖体系。","当：确认补充。","那么：阻止提交并提示至少选择一个覆盖体系，不产生文件记录。"),
        ("场景 7：上传格式/扫描失败","给定：文件超限或病毒扫描失败。","当：上传。","那么：逐文件提示失败，失败文件不计完整度，其他成功文件按既定事务策略处理。"),
        ("场景 8：资料库选择成功","给定：资料库存在有效 Q/E 文件且用户有权限。","当：选择文件并确认。","那么：建立引用、显示来源为资料库选择、更新覆盖与完整度。"),
        ("场景 9：资料库无结果","给定：筛选条件无匹配资料。","当：执行筛选。","那么：显示“当前筛选条件下暂无资料”，确认按钮禁用。"),
        ("场景 10：预览成功","给定：文件存在且用户有查看权限。","当：点击预览。","那么：在受控预览器中打开，记录访问审计，不暴露永久存储地址。"),
        ("场景 11：删除成功","给定：草稿任务中存在可删除文件。","当：在自定义弹窗确认删除。","那么：文件/关联删除，完整性重新计算，toast 明确反馈。"),
        ("场景 12：删除导致缺失","给定：该文件是某体系最后一个有效覆盖。","当：确认删除。","那么：材料状态变为部分覆盖/缺失，对应体系低于100%，提交按钮禁用并说明原因。"),
        ("场景 13：缺失材料禁止提交","给定：F 体系仍缺一个适用必交材料。","当：用户查看提交控制台。","那么：提交按钮禁用，明确列出 F 及缺失材料；构造请求也被服务端拒绝。"),
        ("场景 14：完整材料提交成功","给定：Q/E/S/F 各适用必交材料均为100%，文件均有效。","当：确认提交初次评定。","那么：仅发生一次状态流转，任务进入初次评定/未接收，页面转只读并创建待办。"),
        ("场景 15：重复提交","给定：第一次响应丢失，客户端使用同一幂等键重试。","当：第二次请求到达。","那么：返回相同结果，不重复创建待办和日志。"),
        ("场景 16：并发状态变化","给定：用户编辑时任务被另一端提交或接收。","当：继续上传、删除或提交。","那么：服务端拒绝过期版本并刷新最新状态。"),
        ("场景 17：权限不足","给定：用户无任务或资料库权限。","当：通过 URL/接口访问。","那么：不返回敏感任务/文件信息，编辑请求被拒绝。"),
        ("场景 18：已完结只读","给定：任务已完结。","当：进入详情。","那么：可查看材料、来源和历史，不显示上传、资料库、删除、保存或提交。"),
        ("场景 19：统计失败","给定：完整性服务不可用。","当：页面加载或准备提交。","那么：显示计算失败和重试，提交禁用，不使用旧缓存放行。"),
        ("场景 20：二评退回整改","给定：任务因二次评定退回且列出受影响材料项。","当：项目人员补充并提交整改。","那么：进入二评整改待初评确认，不再经过普通初次评定流程。"),
    ]
    for title,g,w,t in scenarios:
        heading(d,title,3); bullet(d,g); bullet(d,w); bullet(d,t)

    heading(d,"7. 非功能与审计要求",1)
    add_table(d,["类别","要求"],[
        ("性能","48 项材料首屏建议 ≤2秒；筛选和视图切换即时反馈；上传展示逐文件进度。"),
        ("一致性","目录、权限、状态、文件有效性和完整性以后端为准；前端仅提前校验。"),
        ("安全","上传扫描、类型/大小限制、受控预览与下载；所有接口服务端鉴权。"),
        ("审计","上传、资料库关联、删除、提交记录操作人、时间、来源、材料项、体系、任务版本和结果。"),
        ("可访问性","键盘可操作筛选、表格、弹窗和指引；弹窗焦点锁定并可恢复。"),
        ("可追溯","材料项关联文件来源和历史；评定问题关联材料项，不以具体文件形成评定结论。"),
        ("文案","禁用提交时说明缺失体系/材料和下一步；不使用浏览器默认 alert/confirm。"),
    ],[1600,7760],LIGHT_BLUE)

    heading(d,"附录 A：体系与核心术语",1)
    add_table(d,["术语","定义"],[
        ("Q","质量管理体系"),("E","环境管理体系"),("S","职业健康安全"),("F","食品安全体系"),
        ("材料项","MSF11-20 中的业务归档项，是完整性与问题登记的基本对象。"),
        ("具体文件","材料项下的附件证据或在线表单记录，可预览和追溯，但不单独形成评定结论。"),
        ("部分覆盖","已有有效文件，但尚未覆盖该材料在当前项目中的全部适用体系。"),
        ("适用必交材料","结合项目体系、认证类型和清单 requirement 规则判定的提交必备材料。"),
    ],[1800,7560])
    props=d.core_properties
    props.title="归档材料提交详情页功能描述"
    props.subject="体系认证平台归档材料提交详情页功能说明与验收标准"
    props.author="体系认证项目组"
    d.save(OUT); print(OUT)


if __name__ == "__main__":
    build()
